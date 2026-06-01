from collections import defaultdict
import os

from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx2pdf


checkbox='☐'  # Unicode character for an empty checkbox

def create_supplementary_document(date,load,master_bol,pallet_count,dcInfo):
    doc = Document()
    set_margins(doc.sections[0], top=.25, bottom=.25, left=.25, right=.25)
    doc.styles['Normal'].font.name = 'Arial'
    table=doc.add_table(rows=7, cols=1, style='Table Grid')
    table.autofit = False
    #Row 1
    row1 = table.rows[0]  # Clear the default content of the row
    row1.height = Pt(13.5)
    row1_table=row1.cells[0].add_table(rows=1, cols=3)
    row1_table.autofit = True

    clear_space(table)  # Clear space after paragraphs in the row

    addtext('Date:'+date,font_size=9.5,cell=row1_table.cell(0,0))
    row1_table.cell(0,0).width = Inches(1.5)
    addtext('SUPPLEMENT TO THE BILL OF LADING',font_size=13.5,bold=True,cell=row1_table.cell(0,1),align="center")
    row1_table.cell(0,1).width = Inches(5)
    addtext('Page 1',font_size=11.5,cell=row1_table.cell(0,2),bold=True)
    row1_table.cell(0,2).width = Inches(1.5)

    clear_cell(row1.cells[0])  # Clear the default content of the cell
    clear_space(row1_table)  # Clear space after paragraphs in the row

    #Row 2
    row2 = table.rows[1]
    row2.height = Pt(11)
    row2_table=row2.cells[0].add_table(rows=1, cols=3)
    row2_table.autofit = True
    addtext('Bill Lading Number: '+master_bol,font_size=11,cell=row2_table.cell(0,2),bold=True,align="center")
    row2_table.cell(0,0).width = Inches(2.25)
    row2_table.cell(0,1).width = Inches(2.25)
    row2_table.cell(0,2).width = Inches(4)

    clear_cell(row2.cells[0])  # Clear the default content of the cell
    clear_space(row2_table)  # Clear space after paragraphs in the row

    #Row 3
    row3 = table.rows[2]
    row3.height = Pt(8)
    row3._tr.trPr.trHeight.set(qn('w:hRule'), 'exact')

    #Row 4
    row4 = table.rows[3]
    row4.height = Pt(9.5)
    row4._tr.trPr.trHeight.set(qn('w:hRule'), 'exact')

    #Row 5 - PO Info Table
    row5 = table.rows[4]
    row5_table=row5.cells[0].add_table(rows=34, cols=6)
    row5_table.style=doc.styles['Table Grid']
    set_col_widths(row5_table,(Inches(2.25), Inches(.9), Inches(.9), Inches(.6), Inches(.6), Inches(3)))
    row5_table.cell(0,3).merge(row5_table.cell(0,4))

    for i,row in enumerate(row5_table.rows):
        if i==0:
            pass
        else:
            row.height = Pt(15)
            row._tr.trPr.trHeight.set(qn('w:hRule'), 'exact')

    addtext('CUSTOMER ORDER NUMBER',cell=row5_table.cell(0,0),font_size=10.5,bold=True,align="center")
    addtext('# PKGS',cell=row5_table.cell(0,1),font_size=10.5,bold=True,align="center")
    addtext('WEIGHT',cell=row5_table.cell(0,2),font_size=10.5,bold=True,align="center")
    addtext('PALLET/SLIP (CIRCLE ONE)',cell=row5_table.cell(0,3),font_size=8,bold=True,align="center")
    addtext('ADDITIONAL SHIPPER INFO',cell=row5_table.cell(0,5),font_size=10.5,bold=True,align="center")

    weightTotal=0
    qtyTotal=0

    for index,key in enumerate(dcInfo):
        for value in dcInfo[key]:
            qtyTotal+=int(value[1])
            weightTotal+=int(value[2])
            addtext(value[0],font_size=10,cell=row5_table.cell(index+1,0),align="center")
            addtext(value[1],font_size=10,cell=row5_table.cell(index+1,1),align="center")
            addtext(value[2],font_size=10,cell=row5_table.cell(index+1,2),align="center")
            addtext(' ',font_size=10,cell=row5_table.cell(index+1,4),align="center")
            addtext(load,font_size=10,cell=row5_table.cell(index+1,5),align="center")
            index+=1


    clear_cell(row5.cells[0])

    #Row 6
    row6 = table.rows[5]
    row6.height = Pt(9.5)
    row6._tr.trPr.trHeight.set(qn('w:hRule'), 'exact')

    #Row 7
    row7 = table.rows[6]
    row7_table=row7.cells[0].add_table(rows=5, cols=9)
    row7_table.style=doc.styles['Table Grid']
    set_col_widths(row7_table,(Inches(.5), Inches(.5), Inches(.5), Inches(.5), Inches(.9), Inches(.5), Inches(2.75), Inches(.9), Inches(.8)))
    row7_table.cell(0,0).merge(row7_table.cell(0,1))
    row7_table.cell(0,2).merge(row7_table.cell(0,3))
    row7_table.cell(0,7).merge(row7_table.cell(0,8))

    #Row 7-0 Text
    addtext('HANDLING UNIT',cell=row7_table.cell(0,0),font_size=8,bold=True,align="center")
    addtext('PACKAGE',cell=row7_table.cell(0,2),font_size=8.5,bold=True,align="center")
    addtext('QTY',cell=row7_table.cell(0,4),font_size=8.5,bold=True,align="center")
    addtext('COMMODITY DESCRIPTION',cell=row7_table.cell(0,6),font_size=11.5,bold=True,align="center")
    addtext('LTL ONLY',cell=row7_table.cell(0,7),font_size=11.5,bold=True,align="center")

    #Row 7-1 Text
    addtext('QTY',cell=row7_table.cell(1,0),font_size=8.5,bold=True,align="center")
    addtext('TYPE',cell=row7_table.cell(1,1),font_size=8.5,bold=True,align="center")
    addtext('QTY',cell=row7_table.cell(1,2),font_size=8.5,bold=True,align="center")
    addtext('TYPE',cell=row7_table.cell(1,3),font_size=8.5,bold=True,align="center")
    addtext('WEIGHT',cell=row7_table.cell(1,4),font_size=9.5,bold=True,align="center")
    addtext('H.M. (X)',cell=row7_table.cell(1,5),font_size=8,bold=True,align="center")
    addtext('Commodities requiring special or additional care or attention in handling or stowing must be so marked and packaged as to ensure safe transportation with ordinary care. See Section 2(e) of NMFC Item 360',cell=row7_table.cell(1,6),font_size=5,bold=False,align="center")
    addtext('NMFC #',cell=row7_table.cell(1,7),font_size=8.5,bold=True,align="center")
    addtext('CLASS',cell=row7_table.cell(1,8),font_size=8.5,bold=True,align="center")


    #Row 7-2 Text
    addtext(pallet_count,cell=row7_table.cell(2,0),font_size=10,align="center")
    addtext('PLT',cell=row7_table.cell(2,1),font_size=10,align="center")
    addtext(str(qtyTotal),cell=row7_table.cell(2,2),font_size=10,align="center")
    addtext('CTN',cell=row7_table.cell(2,3),font_size=10,align="center")
    addtext(str(weightTotal),cell=row7_table.cell(2,4),font_size=10,align="center")
    addtext('Fashion Jewelry',cell=row7_table.cell(2,6),font_size=10,align="left")
    addtext('100',cell=row7_table.cell(2,8),font_size=10,align="left")

    clear_cell(row7.cells[0]) 
    clear_space(row7_table)

    doc.save('supplementary_document.docx')

def createMasterBOL(master_bol,load_num,date,dc_info,all_bols,client_info,pallet_count):
    doc = Document()
    set_margins(doc.sections[0], top=.25, bottom=.25, left=.25, right=.25)
    doc.styles['Normal'].font.name = 'Arial'
    table=doc.add_table(rows=8, cols=1)
    set_cell_margins(table)
    table.autofit = False

    #Set document borders
    for i,row in enumerate(table.rows):
        if i==0:
            set_cell_border(table.cell(0,0), top={"sz": 10, "val": "single", "color": "000000", "space": "0"})  # Add bottom border to the row

        if i==len(table.rows)-1:
            set_cell_border(table.cell(i,0), bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"})  # Add top border to the row
        
        set_cell_border(table.cell(i,0), start={"sz": 10, "val": "single", "color": "000000", "space": "0"},end={"sz": 12, "val": "single", "color": "000000", "space": "0"})  # Add left border to the row
        #set_cell_border(table.cell(i,1), end={"sz": 12, "val": "single", "color": "000000", "space": "0"})  # Add right border to the row


    #Row 1
    row1 = table.rows[0]  # Clear the default content of the row
    row1.height = Pt(15.5)
    row1_table=row1.cells[0].add_table(rows=1, cols=3)
    row1_table.autofit = True

    clear_space(table)  # Clear space after paragraphs in the row

    addtext('Date:'+date,font_size=9.5,cell=row1_table.cell(0,0))
    row1_table.cell(0,0).width = Inches(1.5)
    addtext('BILL OF LADING',font_size=15.5,bold=True,cell=row1_table.cell(0,1),align="center")
    row1_table.cell(0,1).width = Inches(5)
    addtext('Page 1',font_size=11.5,cell=row1_table.cell(0,2),bold=True)
    row1_table.cell(0,2).width = Inches(1.5)

    clear_cell(row1.cells[0])  # Clear the default content of the cell
    clear_space(row1_table)  # Clear space after paragraphs in the row

    #Row 2
    row2 = table.rows[1]
    row2_table=row2.cells[0].add_table(rows=16, cols=2)
    set_col_widths(row2_table,(Inches(4.5), Inches(3.5)))

    for i,row in enumerate(row2_table.rows):
        set_cell_border(row2_table.cell(i,0), end={"sz": 12, "val": "single", "color": "000000", "space": "0"})

    set_row_height(row2_table.rows[0],Pt(10))
    addtext('SHIP FROM',cell=row2_table.cell(0,0),font_size=8,bold=True,color=RGBColor(255,255,255),align="center",cellBackgroundColor=RGBColor(0,0,0))

    set_cell_border(row2_table.cell(0,1), top={"sz": 10, "val": "single", "color": "000000", "space": "0"})  # Add bottom border to the row


    set_row_height(row2_table.rows[1],Pt(12))
    addtext('Name: ' + client_info[0] + ' / Coast to Coast Fulfillment',cell=row2_table.cell(1,0),font_size=11,bold=True,align="left")
    addtext('Bill of Lading Number:'+master_bol,cell=row2_table.cell(1,1),font_size=11,bold=True,align="left")

    set_row_height(row2_table.rows[2],Pt(12))
    addtext('Address: 773 Victory Highway',cell=row2_table.cell(2,0),font_size=11,bold=True,align="left")
    
    set_row_height(row2_table.rows[3],Pt(12))
    addtext('City/State/Zip: West Greenwich, RI 02817',cell=row2_table.cell(3,0),font_size=11,bold=True,align="left")
    addtext('BAR CODE SPACE',cell=row2_table.cell(3,1),font_size=11,bold=True,align="center",RGBColor=RGBColor(192,192,192))

    clear_cell(row2_table.rows[4].cells[0])  # Clear the default content of the cell
    SID_FOB_inner_table=row2_table.cell(4,0).add_table(rows=1,cols=2)
    set_col_widths(SID_FOB_inner_table,(Inches(3.8), Inches(1.25)))
    addtext('SID#:',cell=SID_FOB_inner_table.cell(0,0),font_size=9,bold=True,align="left")
    addtext('FOB ☐',cell=SID_FOB_inner_table.cell(0,1),font_size=9,bold=True,align="left")
    clear_space(SID_FOB_inner_table)  # Clear space after paragraphs in the row

    set_cell_margins(SID_FOB_inner_table, left=0, right=0)

    set_cell_border(row2_table.cell(4,1), bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    addtext('SHIP TO',cell=row2_table.cell(5,0),font_size=8,bold=True,color=RGBColor(255,255,255),align="center",cellBackgroundColor=RGBColor(0,0,0))
    addtext('CARRIER NAME: '+client_info[7],cell=row2_table.cell(5,1),font_size=11,bold=True,align="left")

    set_row_height(row2_table.rows[6],Pt(40))
    clear_cell(row2_table.rows[6].cells[0])  # Clear the default content of the cell
    ship_to_inner_table=row2_table.cell(6,0).add_table(rows=1,cols=2)
    set_col_widths(ship_to_inner_table,(Inches(2.8), Inches(1.25)))
    addtext('Name: ' + client_info[1],cell=ship_to_inner_table.cell(0,0),font_size=11,bold=True,align="left")
    addtext('Location #:',cell=ship_to_inner_table.cell(0,1),font_size=9,bold=True,align="left")

    addtext('NFI PRO # '+load_num,cell=row2_table.cell(6,1),font_size=11,bold=True,align="left",valign="bottom")
    clear_space(ship_to_inner_table)  # Clear space after paragraphs in the row

    set_cell_margins(ship_to_inner_table, left=0, right=0)

    addtext(client_info[2],cell=row2_table.cell(7,0),font_size=11,bold=True,align="left")
    addtext('Seal numbers(s)',cell=row2_table.cell(7,1),font_size=12,bold=True,align="left")

    set_cell_border(row2_table.cell(7,1), bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    addtext(client_info[4]+', ' + client_info[5]+' '+client_info[6],cell=row2_table.cell(8,0),font_size=11,bold=True,align="left")

    clear_cell(row2_table.rows[9].cells[0])  # Clear the default content of the cell
    SID_FOB_inner_table_2=row2_table.cell(9,0).add_table(rows=1,cols=2)
    set_col_widths(SID_FOB_inner_table_2,(Inches(3.8), Inches(1.25)))
    addtext('FOB ☐',cell=SID_FOB_inner_table_2.cell(0,1),font_size=9,bold=True,align="left")
    clear_space(SID_FOB_inner_table_2)  # Clear space after paragraphs in the row

    set_cell_margins(SID_FOB_inner_table_2, left=0, right=0)

    addtext('Pro number',cell=row2_table.cell(9,1),font_size=9.5,bold=True,align="left")

    addtext('THIRD PART FREIGHT CHARGES BILL TO:',cell=row2_table.cell(10,0),font_size=8,bold=True,color=RGBColor(255,255,255),align="center",cellBackgroundColor=RGBColor(0,0,0))
    addtext('Name:',cell=row2_table.cell(11,0),font_size=9.5,align="left")
    addtext('Address:',cell=row2_table.cell(12,0),font_size=9.5,align="left")

    set_cell_border(row2_table.cell(12,1), bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    addtext('City/State/Zip:',cell=row2_table.cell(13,0),font_size=9.5,align="left")
    set_cell_border(row2_table.cell(13,0), bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    addtext('Freight Charge Terms:',cell=row2_table.cell(13,1),font_size=9.5,bold=True,align="left")
    addtext('(freight charges are prepaid unless marked otherwise)',cell=row2_table.cell(13,1),font_size=8,align="left")

    addtext("Underlying/Sub Bill of Lading #'s\n",cell=row2_table.cell(14,0),font_size=10.5,bold=True,color=RGBColor(255,0,0),align="left")
    addtext(all_bols,cell=row2_table.cell(14,0),font_size=10.5,align="left",bold=True,font_name='Times New Roman')

    set_cell_border(row2_table.cell(14,1), bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    clear_cell(row2_table.cell(14,1))  # Clear the default content of the cell
    Freight_Charge_inner_table=row2_table.cell(14,1).add_table(rows=1,cols=3)
    set_col_widths(Freight_Charge_inner_table,(Inches(1), Inches(1), Inches(1)))
    addtext('Prepaid',cell=Freight_Charge_inner_table.cell(0,0),font_size=9,bold=True,align="left")
    addtext('Collect  X',cell=Freight_Charge_inner_table.cell(0,1),font_size=9,bold=True,align="left")
    addtext('3rd Party',cell=Freight_Charge_inner_table.cell(0,2),font_size=9,bold=True,align="left")
    clear_space(Freight_Charge_inner_table)  # Clear space after paragraphs

    clear_cell(row2_table.cell(15,1))  # Clear the default content of the cell
    master_bol_table=row2_table.cell(15,1).add_table(rows=1,cols=2)
    set_col_widths(master_bol_table,(Inches(1.25), Inches(2)))
    addtext('☐\n',cell=master_bol_table.cell(0,0),font_size=11.5,bold=True,align="center")
    addtext('(check box)',cell=master_bol_table.cell(0,0),font_size=6,bold=True,align="center")
    addtext('Master Bill of Lading:with attached underlying Bills of Ladings',cell=master_bol_table.cell(0,1),font_size=8.5,align="left")
    clear_space(master_bol_table)  # Clear space after paragraphs

    clear_space(row2_table)
    clear_cell(row2.cells[0])  # Clear the default content of the cell

    #Row 3
    addtext('CUSTOMER ORDER INFORMATION',cell=table.cell(2,0),bold=True,font_size=8,align="center",color=RGBColor(255,255,255),cellBackgroundColor=RGBColor(0,0,0))

    #Row 4
    row4 = table.rows[3]
    row4_table=row4.cells[0].add_table(rows=5, cols=6)
    row4_table.style=doc.styles['Table Grid']

    set_col_widths(row4_table,(Inches(2), Inches(1), Inches(1), Inches(.6), Inches(.6), Inches(3)))

    row4_table.cell(0,3).merge(row4_table.cell(0,4))

    for i,row in enumerate(row4_table.rows):
        if i==0:
            pass
        else:
            row.height = Pt(12)
            row._tr.trPr.trHeight.set(qn('w:hRule'), 'exact')

    addtext('CUSTOMER ORDER NUMBER',cell=row4_table.cell(0,0),font_size=8.5,bold=True,align="center")
    addtext('# PKGS',cell=row4_table.cell(0,1),font_size=8.5,bold=True,align="center")
    addtext('WEIGHT',cell=row4_table.cell(0,2),font_size=8.5,bold=True,align="center")
    addtext('PALLET/SLIP\n(CIRCLE ONE)',cell=row4_table.cell(0,3),font_size=8,bold=True,align="center")
    addtext('ADDITIONAL SHIPPER INFO',cell=row4_table.cell(0,5),font_size=8.5,bold=True,align="center")

    po_dic=convertBOLsToDic(dc_info)

    weightTotal=0
    qtyTotal=0

    for index, key in enumerate(po_dic):
        index+=1
        weightTotal+=getPOWeight(po_dic[key])
        qtyTotal+=getPOQty(po_dic[key])
        addtext(key,font_size=10,cell=row4_table.cell(index,0),align="center",bold=True)
        addtext(str(getPOQty(po_dic[key])),font_size=10,cell=row4_table.cell(index,1),align="center",bold=True)
        addtext(str(getPOWeight(po_dic[key])),font_size=10,cell=row4_table.cell(index,2),align="center",bold=True)


    addtext('GRAND TOTAL',cell=row4_table.cell(4,0),font_size=9.5,align="left",bold=True)
    addtext(str(qtyTotal),font_size=9.5,align="center",cell=row4_table.cell(4,1),bold=True)
    addtext(str(weightTotal),font_size=9.5,align="center",cell=row4_table.cell(4,2),bold=True)

    clear_space(row4_table)
    clear_cell(row4.cells[0])  # Clear the default content of the cell

    #Row 5
    addtext('CARRIER INFORMATION',cell=table.cell(4,0),bold=True,font_size=8,align="center",color=RGBColor(255,255,255),cellBackgroundColor=RGBColor(0,0,0))

    #Row 6
    row6 = table.rows[5]
    row6_table=row6.cells[0].add_table(rows=5, cols=9)
    row6_table.style=doc.styles['Table Grid']

    set_col_widths(row6_table,(Inches(.5), Inches(.5), Inches(.5), Inches(.5), Inches(.9), Inches(.5), Inches(2.75), Inches(.9), Inches(1)))

    row6_table.cell(0,0).merge(row6_table.cell(0,1))
    row6_table.cell(0,2).merge(row6_table.cell(0,3))
    row6_table.cell(0,7).merge(row6_table.cell(0,8))

    addtext('HANDLING UNIT',cell=row6_table.cell(0,0),font_size=7.5,bold=True,align="center")
    addtext('PACKAGE',cell=row6_table.cell(0,2),font_size=7.5,bold=True,align="center")
    addtext('COMMODITY DESCRIPTION',cell=row6_table.cell(0,6),font_size=7.5,bold=True,align="center")
    addtext('LTL ONLY',cell=row6_table.cell(0,7),font_size=7.5,bold=True,align="center")

    addtext('QTY',cell=row6_table.cell(1,0),font_size=8.5,bold=True,align="center")
    addtext('TYPE',cell=row6_table.cell(1,1),font_size=8.5,bold=True,align="center")
    addtext('QTY',cell=row6_table.cell(1,2),font_size=8.5,bold=True,align="center")
    addtext('TYPE',cell=row6_table.cell(1,3),font_size=8.5,bold=True,align="center")
    addtext('WEIGHT',cell=row6_table.cell(1,4),font_size=8.5,bold=True,align="center")
    addtext('H.M. (X)',cell=row6_table.cell(1,5),font_size=8,bold=True,align="center")
    addtext('Commodities requiring special or additional care or attention in handling or stowing must be so marked and packaged as to ensure safe transportation with ordinary care. See Section 2(e) of NMFC Item 360',cell=row6_table.cell(1,6),font_size=5,bold=False,align="center")
    addtext('NMFC #',cell=row6_table.cell(1,7),font_size=8.5,bold=True,align="center")
    addtext('CLASS',cell=row6_table.cell(1,8),font_size=8.5,bold=True,align="center")
    addtext(pallet_count,cell=row6_table.cell(2,0),font_size=12,align="center")
    addtext('PLT',cell=row6_table.cell(2,1),font_size=12,align="center")
    addtext(str(qtyTotal),cell=row6_table.cell(2,2),font_size=10,align="center")
    addtext('CTN',cell=row6_table.cell(2,3),font_size=12,align="center")
    addtext(str(weightTotal),cell=row6_table.cell(2,4),font_size=10,align="center")
    addtext(client_info[8],cell=row6_table.cell(2,6),font_size=12,align="left")
    addtext('100',cell=row6_table.cell(2,8),font_size=12,align="left")

    addtext('GRAND TOTAL',font_size=12,cell=row6_table.cell(4,6),align="center",bold=True)

    clear_space(row6_table)
    clear_cell(row6.cells[0])  # Clear the default content of the cell

    #Row 7
    row7 = table.rows[6]
    row7_table=row7.cells[0].add_table(rows=4, cols=2)
    row7_table.style=doc.styles['Table Grid']
    set_col_widths(row7_table,(Inches(4.5), Inches(3.5)))

    row7_table.cell(2,0).merge(row7_table.cell(2,1))

    set_cell_border(row7_table.cell(0,0), end={"sz": 16, "val": "single", "color": "000000", "space": "0"},top={"sz": 12, "val": "single", "color": "000000", "space": "0"})
    set_cell_border(row7_table.cell(0,1), end={"sz": 16, "val": "single", "color": "000000", "space": "0"},top={"sz": 12, "val": "single", "color": "000000", "space": "0"})

    set_cell_border(row7_table.cell(1,1), start={"sz": 16, "val": "single", "color": "000000", "space": "0"},end={"sz": 12, "val": "single", "color": "000000", "space": "0"},bottom={"sz": 12, "val": "single", "color": "000000", "space": "0"})

    addtext("Where the rate is dependent on value, shippers are required to state specifically in writing the agreed or declared value of the property as follows:",cell=row7_table.cell(0,0),font_size=6,align="left")
    addtext("COD Amount $___________________",cell=row7_table.cell(0,1),bold=True,font_size=9.5,align="left")

    addtext("The agreed or declared value of the property is specifically stated by the shipper to be not exceeding __________________ per ___________________.",cell=row7_table.cell(1,0),font_size=6,align="left")
    addtext("Fee Terms: Collect: ☐ Prepaid: ☐ \nCustomer check acceptable: ☐",cell=row7_table.cell(1,1),font_size=9.5,bold=True,align="center")

    addtext("NOTE  Liability Limitation for loss or damage in this shipment may be applicable.  See 49 U.S.C.  14706(c)(1)(A) and (B).",cell=row7_table.cell(2,0),bold=True,font_size=8,align="left")

    addtext("RECEIVED, subject to individually determined rates or contracts that have been agreed upon in writing between the carrier and shipper, if applicable, otherwise to the rates, classifications and rules that have been established by the carrier and are available to the shipper, on request, and to all applicable state and federal regulations.",cell=row7_table.cell(3,0),font_size=6,align="left")
    addtext("The carrier shall not make delivery of this shipment without payment of freight and all other lawful charges\n_______________________________________Shipper Signature\n",cell=row7_table.cell(3,1),font_size=7,align="left")

    clear_space(row7_table)
    clear_cell(row7.cells[0])

    #Row 8
    row8 = table.rows[7]
    row8_table=row8.cells[0].add_table(rows=3, cols=4)
    set_col_widths(row8_table,(Inches(2.3), Inches(1), Inches(1.75), Inches(3)))

    set_cell_border(row8_table.cell(0,0), end={"sz": 10, "val": "single", "color": "000000", "space": "0"})
    set_cell_border(row8_table.cell(0,2), end={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    set_cell_border(row8_table.cell(1,0), end={"sz": 10, "val": "single", "color": "000000", "space": "0"})
    set_cell_border(row8_table.cell(1,2), end={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    set_cell_border(row8_table.cell(2,0), end={"sz": 10, "val": "single", "color": "000000", "space": "0"})
    set_cell_border(row8_table.cell(2,2), end={"sz": 10, "val": "single", "color": "000000", "space": "0"})

    addtext('SHIPPER SIGNATURE / DATE',cell=row8_table.cell(0,0),font_size=11,bold=True,align="left")
    addtext('Trailer Loaded:',cell=row8_table.cell(0,1),font_size=8,underline=True,align="left")
    addtext('Freight Counted:',cell=row8_table.cell(0,2),font_size=8,underline=True,align="left")
    addtext('CARRIER SIGNATURE / PICKUP DATE',cell=row8_table.cell(0,3),font_size=12,bold=True,align="left")

    addtext('This is to certify that the above named materials are properly classified, described, packaged, marked and labeled, and are in proper condition for transportation according to the applicable regulations of the DOT.',cell=row8_table.cell(1,0),font_size=5,align="left")
    addtext('☐ ',cell=row8_table.cell(1,1),font_size=11.5,align="left")
    addtext('By Shipper \n',cell=row8_table.cell(1,1),font_size=7,align="left")
    addtext('☐ ',cell=row8_table.cell(1,1),font_size=11.5,align="left")
    addtext('By Driver',cell=row8_table.cell(1,1),font_size=7,align="left")

    addtext('☐ ',cell=row8_table.cell(1,2),font_size=11.5,align="left")
    addtext('By Shipper \n',cell=row8_table.cell(1,2),font_size=7,align="left")
    addtext('☐ ',cell=row8_table.cell(1,2),font_size=11.5,align="left")
    addtext('By Driver/pallets said to contain\n\n',cell=row8_table.cell(1,2),font_size=7,align="left")

    addtext('Carrier acknowledges receipt of packages and required placards.  Carrier certifies emergency response information was made available and/or carrier has the DOT emergency response guidebook or equivalent documentation in the vehicle.\n',cell=row8_table.cell(1,3),font_size=5,align="left")
    addtext('\nX_______________________\n',cell=row8_table.cell(1,3),font_size=14,align="left")

    addtext('☐ ',cell=row8_table.cell(1,2),font_size=11.5,align="left")
    addtext('By Driver/Pieces',cell=row8_table.cell(1,2),font_size=8,align="left")

    addtext('Property described above is received in good order, except as noted.',cell=row8_table.cell(2,3),font_size=5,align="left")

    clear_space(row8_table)
    clear_cell(row8.cells[0])

    doc.save('master_bol_temp.docx')
    return 'master_bol_temp.pdf'

def set_margins(section, top, bottom, left, right):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def clear_cell(cell):
    for paragraph in cell.paragraphs:
        p=paragraph._element
        p.getparent().remove(p)

def clear_space(table):
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

def addtext(text, font_size=12, bold=False,underline=False,cell=None,color=None,align=None,line_spacing=Pt(0),RGBColor=None,cellBackgroundColor=None,font_name=None,valign=None):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)

    paragraph.paragraph_format.line_spacing = 1.0

    run.font.size = Pt(font_size)
    run.bold = bold
    run.underline = underline
    if color:
        run.font.color.rgb = color

    if RGBColor:
        run.font.color.rgb = RGBColor

    if cellBackgroundColor:
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cellBackgroundColor))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    if font_name:
        run.font.name = font_name

    if align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if valign == "top":
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    elif valign == "center":
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    elif valign == "bottom":
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

def set_col_widths(table,widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def set_row_height(row, height):
    row.height = height
    row._tr.trPr.trHeight.set(qn('w:hRule'), 'exact')

def set_cell_margins(table, left=0, right=0):

    tc = table._element
    tblPr = tc.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    kwargs = {"left":left, "right":right}
    for m in ["left","right"]:
        node = OxmlElement("w:{}".format(m))
        node.set(qn('w:w'), str(kwargs.get(m)))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)

    tblPr.append(tblCellMar)

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def convertBOLsToDic(bolList):
    bolDic=defaultdict(list)
    for bol in bolList:
        for subArray in bol[3]:
            po_num=subArray[0].rsplit("-",1)[0]
            bolDic[po_num].append(subArray)
    return bolDic

def getPOWeight(poInfo):
    weight=0
    for subArray in poInfo:
        weight+=int(subArray[2])
    return weight

def getPOQty(poInfo):
    qty=0
    for subArray in poInfo:
        qty+=int(subArray[1])
    return qty